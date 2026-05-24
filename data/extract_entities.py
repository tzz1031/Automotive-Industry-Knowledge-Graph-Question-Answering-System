import json
import os
import re
import time
import argparse
from collections import defaultdict
from pathlib import Path

BASE_DIR = Path(__file__).resolve().parent.parent
TECH_DEFAULT_OUTPUT = BASE_DIR / "data" / "car_data" / "tech_entity_centric.json"
TECH_DEFAULT_SOURCE_INDEX = BASE_DIR / "data" / "car_data" / "knowledge_sources.json"

# ==================== 配置区 ====================
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "sk-e328faa7012c4272b1d6057efc24dee5")
DEEPSEEK_BASE_URL = "https://api.deepseek.com"
MODEL = "deepseek-chat"
CHUNK_SIZE = 3000          # 每块字符数（增大以容纳更多内容）
MAX_RETRIES = 3            # API 重试次数
INPUT_FILE = "demo.txt"
OUTPUT_FILE = "entity_centric.json"
# =================================================

SYSTEM_PROMPT = """你是一个中文知识图谱三元组提取专家。请从给定文本中提取所有事实信息，输出 JSON。

## 输出格式
{
  "triples": [
    {
      "subject": "实体或主题",
      "relation": "关系或属性名",
      "object": "实体或属性值",
      "type": "attribute / relation"
    }
  ]
}

## type 判断规则
1. **type="attribute"**：subject 的属性或统计数值
   - 产销数据："中国汽车 销量 3143.6万辆"
   - 同比变化："比亚迪 同比增长 41.1%"
   - 占比份额："新能源 渗透率 40.9%"
   - 排名位次："比亚迪 排名 第一"
   - 技术规格："神盾短刀电池 能量密度 192Wh/kg"
   - 产品属性："蔚来ET9 搭载 ADS 3.0"

2. **type="relation"**：两个实体之间的关系
   - 竞争："比亚迪 超越 一汽大众"
   - 发布："宁德时代 发布 骁遥超级增混电池"
   - 搭载："问界M7 搭载 192线激光雷达"
   - 出口："中国 新能源汽车出口 比利时" (需同时提取数量为 attribute)
   - 合作："华为 合作 科大讯飞"

## 提取规则
1. relation 要**简洁规范**，如用"发布"而非"于2024年4月隆重发布"
2. object 值**保留原文中的数字和单位**（万辆、亿元、%、Wh/kg、km、TOPS、mAh、kW 等）
3. **不遗漏**：每个统计数字都要提取，包括年份、百分比、排名、变化趋势
4. **主语统一**：指代中国/我国统一用"中国"；企业用全称；产品名用原文完整名称
5. **技术产品**：每个产品/技术发布会提取"X发布Y"的关系和"Y的规格"的属性
6. **出口国家**：对每个国家的出口数据单独提取"""

ENTITY_ALIAS = {
    # 宏观主体
    "中国汽车市场": "中国", "中国汽车": "中国", "我国": "中国",
    "中国品牌": "中国品牌", "中国车企": "中国车企",
    # 车企归一化
    "奇瑞集团": "奇瑞", "吉利汽车": "吉利", "吉利控股集团": "吉利",
    "一汽-大众": "一汽大众", "一汽大众": "一汽大众",
    "上海汽车集团股份有限公司": "上汽集团",
    "上汽大众": "上汽大众", "一汽奥迪": "一汽奥迪", "一汽": "中国一汽",
    "华晨宝马": "华晨宝马", "北京奔驰": "北京奔驰",
    "东风本田": "东风本田", "广汽本田": "广汽本田",
    "广汽丰田": "广汽丰田", "一汽丰田": "一汽丰田",
    "东风日产": "东风日产", "上汽通用": "上汽通用",
    "理想汽车": "理想汽车", "长安汽车": "长安", "长城汽车": "长城",
    "蔚来汽车": "蔚来", "广汽埃安": "广汽埃安", "深蓝汽车": "深蓝",
    "腾势": "腾势", "极狐汽车": "极狐", "岚图汽车": "岚图",
    "小米汽车": "小米汽车", "福田汽车": "福田",
    # 供应商
    "宁德时代": "宁德时代", "蜂巢能源": "蜂巢能源", "欣旺达": "欣旺达",
    "赣锋锂电": "赣锋锂电", "国轩高科": "国轩高科", "汇川联合动力": "汇川联合动力",
    "威睿": "威睿", "星驱科技": "星驱科技", "上海电驱动": "上海电驱动",
    "马瑞利": "马瑞利", "速豹科技": "速豹科技",
    # 智驾/芯片
    "华为": "华为", "百度": "百度", "百度Apollo": "百度",
    "Momenta": "Momenta", "科大讯飞": "科大讯飞",
    "速腾聚创": "速腾聚创", "赛恩领动": "赛恩领动",
    "联发科": "联发科", "芯驰科技": "芯驰科技", "为旌科技": "为旌科技",
    "凯芯科技": "凯芯科技", "宸芯科技": "宸芯科技",
    "安波福": "安波福", "英飞源": "英飞源",
    "优优绿能": "优优绿能", "绿能慧充": "绿能慧充",
    "育材堂": "育材堂", "安徽宝镁": "安徽宝镁",
    "智己": "智己", "小米": "小米",
    # 车型/品类归一化
    "新能源乘用车": "新能源乘用车", "新能源商用车": "新能源商用车",
    "纯电动车": "纯电动汽车", "纯电动汽车": "纯电动汽车",
    "插电式混合动力汽车": "插混汽车", "PHEV（含REEV）": "插混汽车",
    "插混": "插混汽车", "纯电": "纯电动汽车",
    "燃料电池汽车": "燃料电池汽车",
    "乘用车": "乘用车", "商用车": "商用车", "新能源汽车": "新能源汽车",
    # 概念归一化
    "前十大汽车集团": "全球TOP10汽车集团",
    "TOP10企业": "中国TOP10企业", "TOP10出口车企": "中国出口TOP10",
    "新能源汽车TOP10企业": "新能源汽车TOP10",
    "进口汽车": "进口汽车", "传统燃油车巨头": "传统燃油车企",
}

FUZZY_OBJECTS = {
    "是", "否", "下降", "下滑", "萎缩", "增长", "低迷", "被挤压",
    "销量下滑", "销量跃升", "全球冠军", "全球销量前十",
    "全球新能源汽车领域领先地位", "多品牌战略", "全球化布局",
    "主要细分市场", "第一", "第二", "第三",
}

def read_input(filepath: str) -> str:
    """读取文件，自动识别 .docx 或 .txt"""
    if filepath.endswith(".docx"):
        from docx import Document
        doc = Document(filepath)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    else:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read().strip()


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE) -> list[str]:
    """按空行切块，块内保持段落完整，不超过 chunk_size"""
    paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
    chunks = []
    current = ""

    for para in paragraphs:
        if not current:
            current = para
        elif len(current) + len(para) + 1 <= chunk_size:
            current = current + "\n" + para
        else:
            chunks.append(current)
            if len(para) > chunk_size:
                chunks.extend(_split_long(para, chunk_size))
                current = ""
            else:
                current = para

    if current:
        chunks.append(current)
    return chunks


def _split_long(text: str, chunk_size: int) -> list[str]:
    sentences = text.replace("；", "。").replace("。", "。\n").split("\n")
    chunks, current = [], ""
    for s in sentences:
        s = s.strip()
        if not s: continue
        if len(current) + len(s) + 1 <= chunk_size:
            current = (current + "\n" + s).strip() if current else s
        else:
            if current: chunks.append(current)
            current = s
    if current: chunks.append(current)
    return chunks


def _extract_from_chunk(client, text: str) -> list[dict]:
    for attempt in range(MAX_RETRIES):
        try:
            response = client.chat.completions.create(
                model=MODEL,
                messages=[
                    {"role": "system", "content": SYSTEM_PROMPT},
                    {"role": "user", "content": f"请从以下文本中提取所有三元组：\n\n{text}"}
                ],
                response_format={"type": "json_object"},
                temperature=0.1,
                max_tokens=8192,
            )
            content = response.choices[0].message.content
            return json.loads(content).get("triples", [])
        except json.JSONDecodeError:
            print(f"  [WARN] JSON解析失败(尝试{attempt+1}): {content[:80]}...")
            if attempt < MAX_RETRIES - 1: time.sleep(2)
        except Exception as e:
            print(f"  [WARN] API错误(尝试{attempt+1}): {e}")
            if attempt < MAX_RETRIES - 1: time.sleep(3)
    return []


def extract_triples(filepath: str) -> list[dict]:
    print(f"[1/3] 读取: {filepath}")
    text = read_input(filepath)
    print(f"  总字符数: {len(text)}")

    chunks = chunk_text(text)
    print(f"  分成 {len(chunks)} 个 chunk (每块 ≤ {CHUNK_SIZE} 字符)")

    from openai import OpenAI
    client = OpenAI(api_key=DEEPSEEK_API_KEY, base_url=DEEPSEEK_BASE_URL)
    all_triples = []
    for i, chunk in enumerate(chunks, 1):
        print(f"  [{i}/{len(chunks)}] ({len(chunk)} 字符) ...", end=" ")
        triples = _extract_from_chunk(client, chunk)
        all_triples.extend(triples)
        print(f"→ {len(triples)} 条")

    # 简单去重
    seen = set()
    unique = []
    for t in all_triples:
        key = (t.get("subject",""), t.get("relation",""), t.get("object",""))
        if key not in seen:
            seen.add(key); unique.append(t)

    print(f"  原始 {len(all_triples)} 条, 去重后 {len(unique)} 条\n")
    return unique


# ==================== 后处理 ====================

def _normalize_triple(t: dict) -> dict:
    return {
        "subject": ENTITY_ALIAS.get(t["subject"], t["subject"]),
        "relation": t["relation"],
        "object": ENTITY_ALIAS.get(t["object"], t["object"]),
        "type": t.get("type", "attribute"),
    }


def _is_valid(t: dict) -> bool:
    if t["object"] in FUZZY_OBJECTS: return False
    if len(t["object"]) > 80: return False
    if not t["subject"] or not t["object"]: return False
    if re.match(r"^(超过|大于|高于|突破)\d+", t["object"]): return False
    return True


def _norm_val(val: str) -> str:
    return re.sub(r"\s+", "", str(val).strip())


def _relation_overlap(a: str, b: str) -> bool:
    if a in b or b in a: return True
    a1 = re.sub(r"2024年|2023年|2025年|较上年|同比|全年|合计", "", a)
    b1 = re.sub(r"2024年|2023年|2025年|较上年|同比|全年|合计", "", b)
    return a1 == b1 and bool(a1)


def _simplify_relation(r: str) -> str:
    r = re.sub(r"^(2024年|2025年|全年|累计)", "", r)
    r = r.strip()
    map_replace = {
        "向全球供应电池材料占比": "电池材料全球供应占比",
        "向全球供应动力电池占比": "动力电池全球供应占比",
        "占亚洲总销量比重": "亚洲市场份额",
        "占全球销量比重": "全球市场份额",
        "占全年总销量比重": "国内市场份额",
        "占出口总量比重": "出口份额",
        "新能源汽车销量占汽车总销量比重": "新能源渗透率",
        "出口较上年同期占比": "出口占比变动",
    }
    r = map_replace.get(r, r)
    r = re.sub(r"([产销]量)(同比[增减])", r"\1\2", r)
    return re.sub(r"\s+", "", r)


def _deduplicate_deep(triples: list[dict]) -> list[dict]:
    seen = {}
    for t in triples:
        key = (t["subject"], t["relation"], _norm_val(t["object"]))
        if key not in seen or len(t["relation"]) < len(seen[key]["relation"]):
            seen[key] = t

    by_key = defaultdict(list)
    for t in seen.values():
        by_key[(t["subject"], _norm_val(t["object"]))].append(t)

    result = []
    for group in by_key.values():
        if len(group) == 1:
            result.append(group[0])
        else:
            kept = {}
            for t in group:
                if not any(_relation_overlap(t["relation"], kr) for kr in kept):
                    kept[t["relation"]] = t
            result.extend(kept.values())
    return result


def post_process(triples: list[dict]) -> tuple[list, list]:
    print("[2/3] 后处理...")
    triples = [_normalize_triple(t) for t in triples]
    before = len(triples)
    triples = [t for t in triples if _is_valid(t)]
    print(f"  过滤: {before} → {len(triples)} (-{before-len(triples)})")
    before = len(triples)
    triples = _deduplicate_deep(triples)
    print(f"  去重: {before} → {len(triples)}")
    for t in triples:
        t["relation"] = _simplify_relation(t["relation"])
    before = len(triples)
    triples = _deduplicate_deep(triples)
    print(f"  简化后去重: {before} → {len(triples)}")

    attrs = [t for t in triples if t["type"] == "attribute"]
    rels = [t for t in triples if t["type"] == "relation"]
    print(f"  属性 {len(attrs)} + 关系 {len(rels)} = {len(triples)}\n")
    return attrs, rels


# ==================== 实体中心化 ====================

def _to_number(val):
    s = str(val)
    for unit in ["万亿元", "万亿元", "亿元", "万辆", "个百分点", "个位次",
                 "Wh/kg", "Wh/L", "TOPS", "mAh", "kW", "km", "%",
                 "倍", "辆", "个"]:
        if s.endswith(unit):
            s = s[:-len(unit)]
            break
    s = s.replace("超过", "").replace("突破", "").replace("微降", "-")
    s = s.replace("微增", "+").replace("约", "").replace("大于", "").replace("连续", "")
    s = s.strip()
    try:
        return int(s) if float(s) == int(float(s)) else float(s)
    except (ValueError, TypeError):
        return None


def build_entity_graph(attrs: list[dict], rels: list[dict]) -> dict:
    entities = {}

    for t in attrs:
        subj = t["subject"]
        if subj not in entities:
            entities[subj] = {"type": "entity", "attributes": {}, "relations": []}
        num = _to_number(t["object"])
        entities[subj]["attributes"][t["relation"]] = num if num is not None else str(t["object"])

    for t in rels:
        subj = t["subject"]
        if subj not in entities:
            entities[subj] = {"type": "entity", "attributes": {}, "relations": []}
        for r in entities[subj]["relations"]:
            if r["relation"] == t["relation"]:
                obj_list = r["object"] if isinstance(r["object"], list) else [r["object"]]
                obj_list.append(t["object"])
                r["object"] = list(dict.fromkeys(obj_list))  # 去重保序
                break
        else:
            entities[subj]["relations"].append({"relation": t["relation"], "object": t["object"]})

    priority = [
        "中国", "比亚迪", "吉利", "奇瑞", "长安", "长城",
        "一汽大众", "上汽集团", "特斯拉", "蔚来", "理想汽车",
        "华为", "宁德时代", "百度", "小米",
        "中国TOP10企业", "全球TOP10汽车集团", "新能源汽车TOP10",
        "中国出口TOP10",
    ]
    ordered = {}
    for name in priority:
        if name in entities:
            ordered[name] = entities.pop(name)
    for name in sorted(entities.keys()):
        ordered[name] = entities[name]
    return ordered


# ==================== 主流程 ====================

def main_legacy():
    raw = extract_triples(INPUT_FILE)
    attrs, rels = post_process(raw)

    print("[3/3] 生成实体中心化 JSON...")
    graph = build_entity_graph(attrs, rels)

    total_attr = sum(len(e["attributes"]) for e in graph.values())
    total_rel = sum(len(e["relations"]) for e in graph.values())

    output = {
        "meta": {
            "source": INPUT_FILE,
            "entities": len(graph),
            "total_attributes": total_attr,
            "total_relations": total_rel,
        },
        "entities": graph,
    }

    with open(OUTPUT_FILE, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f"  已生成 {OUTPUT_FILE}")
    print(f"  实体: {len(graph)}  属性: {total_attr}  关系: {total_rel}")


TECH_SEED_TRIPLES = [
    {"subject": "GB 38031-2025", "relation": "技术类别", "object": "Standard", "type": "attribute", "source": "电动汽车用动力蓄电池安全要求", "evidence": "动力蓄电池安全标准"},
    {"subject": "GB 38031-2025", "relation": "规范对象", "object": "动力蓄电池", "type": "relation", "source": "电动汽车用动力蓄电池安全要求", "evidence": "规定动力蓄电池安全要求"},
    {"subject": "GB/T 31467", "relation": "规范对象", "object": "锂离子动力蓄电池包和系统", "type": "relation", "source": "动力电池包和系统测试规程", "evidence": "定义动力电池包和系统测试规程"},
    {"subject": "GB/T 18487.1", "relation": "规范对象", "object": "传导充电系统", "type": "relation", "source": "电动汽车传导充电系统通用要求", "evidence": "规定传导充电系统通用要求"},
    {"subject": "GB/T 27930", "relation": "规范对象", "object": "非车载充电机与电动汽车通信", "type": "relation", "source": "充电通信协议", "evidence": "定义充电机与电动汽车之间通信协议"},
    {"subject": "麒麟电池", "relation": "技术类别", "object": "BatteryTechnology", "type": "attribute", "source": "宁德时代 CTP 3.0 麒麟电池资料", "evidence": "麒麟电池属于动力电池技术"},
    {"subject": "麒麟电池", "relation": "采用", "object": "CTP 3.0", "type": "relation", "source": "宁德时代 CTP 3.0 麒麟电池资料", "evidence": "麒麟电池基于 CTP 3.0"},
    {"subject": "麒麟电池", "relation": "系统能量密度", "object": "最高255Wh/kg", "type": "attribute", "source": "宁德时代 CTP 3.0 麒麟电池资料", "evidence": "公开资料提到系统能量密度最高255Wh/kg"},
    {"subject": "麒麟电池", "relation": "研发企业", "object": "宁德时代", "type": "relation", "source": "宁德时代 CTP 3.0 麒麟电池资料", "evidence": "宁德时代发布麒麟电池"},
    {"subject": "刀片电池", "relation": "技术类别", "object": "BatteryTechnology", "type": "attribute", "source": "比亚迪刀片电池资料", "evidence": "刀片电池属于动力电池技术"},
    {"subject": "刀片电池", "relation": "正极材料", "object": "磷酸铁锂", "type": "relation", "source": "比亚迪刀片电池资料", "evidence": "刀片电池采用磷酸铁锂体系"},
    {"subject": "刀片电池", "relation": "研发企业", "object": "比亚迪", "type": "relation", "source": "比亚迪刀片电池资料", "evidence": "比亚迪发布刀片电池"},
    {"subject": "800V高压平台", "relation": "技术类别", "object": "ChargingTechnology", "type": "attribute", "source": "高压快充技术资料", "evidence": "800V平台用于高压快充"},
    {"subject": "800V高压平台", "relation": "支持", "object": "超快充", "type": "relation", "source": "高压快充技术资料", "evidence": "800V高压平台提升补能效率"},
    {"subject": "华为MDC", "relation": "技术类别", "object": "ComputingPlatform", "type": "attribute", "source": "华为MDC白皮书", "evidence": "MDC是智能驾驶计算平台"},
    {"subject": "华为MDC", "relation": "研发企业", "object": "华为", "type": "relation", "source": "华为MDC白皮书", "evidence": "华为推出MDC平台"},
    {"subject": "华为ADS", "relation": "技术类别", "object": "ADASSystem", "type": "attribute", "source": "华为智能汽车解决方案资料", "evidence": "ADS是智能驾驶系统"},
    {"subject": "华为ADS", "relation": "包含能力", "object": ["城区领航辅助", "高速领航辅助", "泊车辅助"], "type": "relation", "source": "华为智能汽车解决方案资料", "evidence": "ADS覆盖领航辅助和泊车辅助场景"},
    {"subject": "征程5", "relation": "技术类别", "object": "Chip", "type": "attribute", "source": "地平线征程5资料", "evidence": "征程5是智能驾驶芯片"},
    {"subject": "征程5", "relation": "算力", "object": "128TOPS", "type": "attribute", "source": "地平线征程5资料", "evidence": "征程5公开资料标注128TOPS"},
    {"subject": "征程5", "relation": "研发企业", "object": "地平线", "type": "relation", "source": "地平线征程5资料", "evidence": "地平线发布征程5"},
    {"subject": "NVIDIA DRIVE Orin", "relation": "技术类别", "object": "Chip", "type": "attribute", "source": "NVIDIA DRIVE Orin资料", "evidence": "Orin是自动驾驶SoC"},
    {"subject": "NVIDIA DRIVE Orin", "relation": "典型算力", "object": "最高254TOPS", "type": "attribute", "source": "NVIDIA DRIVE Orin资料", "evidence": "Orin公开资料标注最高254TOPS"},
    {"subject": "NVIDIA DRIVE Orin", "relation": "研发企业", "object": "NVIDIA", "type": "relation", "source": "NVIDIA DRIVE Orin资料", "evidence": "NVIDIA推出DRIVE Orin"},
    {"subject": "骁龙8295", "relation": "技术类别", "object": "Chip", "type": "attribute", "source": "高通骁龙座舱平台资料", "evidence": "8295是智能座舱芯片"},
    {"subject": "骁龙8295", "relation": "支持", "object": ["车机计算", "多屏交互", "OTA升级", "C-V2X"], "type": "relation", "source": "高通骁龙座舱平台资料", "evidence": "8295用于智能座舱计算"},
    {"subject": "禾赛AT128", "relation": "技术类别", "object": "Sensor", "type": "attribute", "source": "禾赛AT128资料", "evidence": "AT128是车规级激光雷达"},
    {"subject": "禾赛AT128", "relation": "线数", "object": "128线", "type": "attribute", "source": "禾赛AT128资料", "evidence": "AT128名称和资料体现128线激光雷达"},
    {"subject": "禾赛AT128", "relation": "研发企业", "object": "禾赛科技", "type": "relation", "source": "禾赛AT128资料", "evidence": "禾赛科技推出AT128"},
    {"subject": "速腾聚创M1", "relation": "技术类别", "object": "Sensor", "type": "attribute", "source": "速腾聚创M1资料", "evidence": "M1是车规级激光雷达"},
    {"subject": "速腾聚创M1", "relation": "技术路线", "object": "MEMS", "type": "attribute", "source": "速腾聚创M1资料", "evidence": "M1采用MEMS路线"},
    {"subject": "C-V2X", "relation": "技术类别", "object": "CommunicationTechnology", "type": "attribute", "source": "车联网技术资料", "evidence": "C-V2X是车路协同通信技术"},
    {"subject": "C-V2X", "relation": "支持", "object": ["车车通信", "车路通信", "车云通信", "车人通信"], "type": "relation", "source": "车联网技术资料", "evidence": "C-V2X支持V2V/V2I/V2N/V2P"},
    {"subject": "雷达-相机融合", "relation": "技术类别", "object": "AlgorithmTechnology", "type": "attribute", "source": "Radar-Camera Fusion Review", "evidence": "雷达-相机融合属于多传感器融合技术"},
    {"subject": "雷达-相机融合", "relation": "融合传感器", "object": ["毫米波雷达", "摄像头"], "type": "relation", "source": "Radar-Camera Fusion Review", "evidence": "融合毫米波雷达和摄像头用于感知"},
]


TECH_ALIAS = {
    "Orin": "NVIDIA DRIVE Orin",
    "DRIVE Orin": "NVIDIA DRIVE Orin",
    "AT128": "禾赛AT128",
    "M1": "速腾聚创M1",
    "Journey 5": "征程5",
    "Qualcomm 8295": "骁龙8295",
    "SA8295P": "骁龙8295",
}


TECH_TERM_CATEGORY = {
    "GEA智能新能源架构": "VehiclePlatform",
    "增程式平台技术": "VehiclePlatform",
    "蔚来ET9数字架构": "EEArchitecture",
    "欧曼银河9重卡平台": "VehiclePlatform",
    "骁遥超级增混电池": "BatteryTechnology",
    "59度大电量增混电池": "BatteryTechnology",
    "软包CTP一体化电池": "BatteryTechnology",
    "神盾短刀磷酸铁锂电池": "BatteryTechnology",
    "L600启晨电池": "BatteryTechnology",
    "准900V超快充固态电池": "BatteryTechnology",
    "金沙江电池": "BatteryTechnology",
    "非晶碳纤维电机": "MotorTechnology",
    "P1电机": "MotorTechnology",
    "800V SiC四电机分布式驱动系统": "DriveSystem",
    "三电机构型": "DriveSystem",
    "超高压双电机电驱系统": "DriveSystem",
    "分布式电驱桥系统": "DriveSystem",
    "第二代增程技术": "DriveSystem",
    "PD4H混碳电控": "PowerElectronics",
    "磐石底盘": "ChassisTechnology",
    "智能底盘预研技术": "ChassisTechnology",
    "天工08底盘": "ChassisTechnology",
    "闪充电池3.0": "ChargingTechnology",
    "720kW独立风道超充系统": "ChargingTechnology",
    "兆瓦级充电堆": "ChargingTechnology",
    "地埋式全液冷充电系统": "ChargingTechnology",
    "OTA 5.0": "SoftwareTechnology",
    "NOP+": "ADASSystem",
    "达尔文2.0技术体系": "TechnologySystem",
    "GARCIA OS": "ADASSystem",
    "智能座舱": "SmartCockpit",
    "192线激光雷达": "Sensor",
    "NOP智慧领航辅助驾驶系统": "ADASSystem",
    "C-Pilot 5.0": "ADASSystem",
    "SIR-4K": "Sensor",
    "M3": "Sensor",
    "ANP3 Pro": "ADASSystem",
    "AD 5.0智驾大模型": "AlgorithmTechnology",
    "第六代ADAS平台": "ADASSystem",
    "讯飞星火大模型V4.0": "AlgorithmTechnology",
    "ADS 3.0": "ADASSystem",
    "御行系列域控芯片": "Chip",
    "VS919": "Chip",
    "VS919L": "Chip",
    "KT5030A": "Chip",
    "Dimensity Auto座舱平台系统单芯片SoC": "Chip",
    "1+N中央计算+区域控制架构": "EEArchitecture",
    "2000MPa激光拼焊门环": "MaterialProcess",
    "镁合金CCB": "MaterialProcess",
    "镁合金电驱壳体": "MaterialProcess",
    "半固态工艺镁合金电驱壳体": "MaterialProcess",
    "智能电池": "BatteryTechnology",
    "自动驾驶运行安全风险管控系统": "SafetySystem",
    "城市NOA": "ADASFunction",
    "端到端自动驾驶": "AlgorithmTechnology",
    "车路云一体化": "V2XTechnology",
}


TECH_COMPANIES = [
    "吉利", "深蓝汽车", "蔚来", "福田汽车", "宁德时代", "蜂巢能源", "赣锋锂电", "国轩高科",
    "智己", "小米", "广汽埃安", "星驱科技", "威睿", "腾势", "马瑞利", "速豹科技", "上海电驱动",
    "汇川联合动力", "优优绿能", "绿能慧充", "英飞源", "理想汽车", "极狐汽车", "广汽集团",
    "岚图汽车", "华为", "上汽通用", "奇瑞", "赛恩领动", "速腾聚创", "百度Apollo", "Momenta",
    "安波福", "科大讯飞", "为旌科技", "凯芯科技", "联发科", "芯驰科技", "育材堂", "安徽宝镁",
    "上汽集团", "比亚迪", "长城", "小鹏", "商汤绝影", "元戎启行", "零一汽车", "清华大学", "一汽",
]


def clean_tech_name(name):
    name = re.sub(r"^[，,。；;：:\s]*(2024年|2025年|[0-9]{1,2}月[0-9]{0,2}日)?", "", str(name))
    name = re.sub(r"^(发布|推出|采用|搭载|支持|实现|具有|具备|该|这一|这些|其|同时|此外|通过|随着)", "", name)
    name = re.sub(r"^(了|的|全新|新一代|面向[^的]{1,20}的)", "", name)
    name = re.sub(r"(升级|技术|系统|方案|平台|产品|芯片|电池|底盘|雷达|电机|电控|架构)(等.*)$", r"\1", name)
    name = re.sub(r"[\s，,。；;：:]+$", "", name)
    return name.strip()


def is_valid_tech_entity(name):
    name = str(name).strip()
    if not (1 < len(name) <= 35):
        return False
    bad_prefixes = ("该", "这一", "这些", "同时", "此外", "通过", "随着", "利用", "结合", "国管局", "标准引领")
    if name.startswith(bad_prefixes):
        return False
    if name.count("、") >= 3:
        return False
    if any(mark in name for mark in ["\n", "。", "；"]):
        return False
    return True


def compact_for_match(text):
    return re.sub(r"[\s“”\"'（）()]+", "", str(text))


def normalize_tech_name(name):
    name = clean_tech_name(re.sub(r"\s+", " ", str(name)).strip())
    return TECH_ALIAS.get(name, name)


def normalize_tech_relation(relation):
    relation = str(relation).strip()
    mapping = {
        "研发公司": "研发企业",
        "开发企业": "研发企业",
        "生产企业": "研发企业",
        "应用": "应用场景",
        "适用": "应用场景",
        "规范": "规范对象",
    }
    return mapping.get(relation, relation)


def read_technical_text(path):
    path = Path(path)
    suffix = path.suffix.lower()
    if suffix in {".txt", ".md", ".json", ".csv"}:
        return path.read_text(encoding="utf-8", errors="ignore")
    if suffix == ".docx":
        from docx import Document
        doc = Document(str(path))
        return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
    if suffix == ".pdf":
        try:
            import fitz
            parts = []
            with fitz.open(str(path)) as doc:
                for page in doc:
                    parts.append(page.get_text("text"))
            return "\n".join(parts)
        except ImportError:
            try:
                import pdfplumber
            except ImportError as exc:
                raise RuntimeError("PDF extraction requires PyMuPDF or pdfplumber.") from exc
            parts = []
            with pdfplumber.open(str(path)) as pdf:
                for page in pdf.pages:
                    parts.append(page.extract_text() or "")
            return "\n".join(parts)
    raise ValueError(f"Unsupported input file type: {path}")


def extract_technical_triples_by_rules(text, source_name):
    triples = []
    hints = {
        "GB 38031": ("技术类别", "Standard"),
        "GB/T 31467": ("规范对象", "锂离子动力蓄电池包和系统"),
        "GB/T 18487": ("规范对象", "传导充电系统"),
        "GB/T 27930": ("规范对象", "充电通信协议"),
        "麒麟电池": ("研发企业", "宁德时代"),
        "刀片电池": ("研发企业", "比亚迪"),
        "800V": ("支持", "超快充"),
        "MDC": ("研发企业", "华为"),
        "ADS": ("研发企业", "华为"),
        "征程5": ("研发企业", "地平线"),
        "Orin": ("研发企业", "NVIDIA"),
        "8295": ("研发企业", "高通"),
        "AT128": ("研发企业", "禾赛科技"),
        "C-V2X": ("支持", "车路协同"),
    }
    compact_text = compact_for_match(text)
    for term, category in TECH_TERM_CATEGORY.items():
        if compact_for_match(term) in compact_text:
            triples.append({"subject": term, "relation": "技术类别", "object": category, "type": "attribute", "source": source_name, "evidence": term})

    sentences = [s.strip() for s in re.split(r"[。！？!?；;\n]+", re.sub(r"\s+", " ", text)) if len(s.strip()) >= 8]
    for sentence in sentences:
        for keyword, (rel, obj) in hints.items():
            if keyword in sentence:
                subject = normalize_tech_name(keyword)
                triples.append({"subject": subject, "relation": rel, "object": obj, "type": "relation", "source": source_name, "evidence": sentence[:240]})

        compact_sentence = compact_for_match(sentence)
        for term, category in TECH_TERM_CATEGORY.items():
            if compact_for_match(term) in compact_sentence:
                triples.append({"subject": term, "relation": "技术类别", "object": category, "type": "attribute", "source": source_name, "evidence": sentence[:240]})
                for company in TECH_COMPANIES:
                    if company in sentence and re.search(r"(发布|推出|展示|宣布|官宣|升级|打造|采用|搭载)", sentence):
                        triples.append({"subject": term, "relation": "研发企业", "object": company, "type": "relation", "source": source_name, "evidence": sentence[:240]})

                value_patterns = [
                    ("算力", r"([0-9.]+\s*TOPS)"),
                    ("能量密度", r"([0-9.]+\s*Wh/kg|[0-9.]+\s*Wh/L)"),
                    ("测距能力", r"([0-9.]+m@[0-9.]+%)"),
                    ("角分辨率", r"([0-9.]+°×[0-9.]+°)"),
                    ("峰值功率", r"([0-9.]+(?:～|-|~)[0-9.]+\s*kW)"),
                    ("单枪最大功率", r"单枪最大支持\s*([0-9.]+\s*kW)"),
                    ("噪声", r"噪声小于\s*([0-9.]+\s*dB)"),
                    ("效率", r"效率\s*([0-9.]+%以上?|大于[0-9.]+%)"),
                    ("充电倍率", r"([0-9.]+C)"),
                    ("输出规格", r"([0-9.]+V/[0-9.]+A)"),
                    ("续航里程", r"([0-9.]+\+?\s*km)"),
                    ("循环寿命", r"([0-9.]+次)"),
                    ("减重比例", r"(?:减重|体积缩小|提升)\s*(超|超过|至|达)?\s*([0-9.]+%|1/3)"),
                ]
                for rel, value_re in value_patterns:
                    for vm in re.finditer(value_re, sentence, flags=re.I):
                        value = vm.group(vm.lastindex or 1)
                        triples.append({"subject": term, "relation": rel, "object": value.replace(" ", ""), "type": "attribute", "source": source_name, "evidence": sentence[:240]})

        patterns = [
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:算力|AI算力)[为可达约最高]*\s*([0-9.]+\s*TOPS)", "算力"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:能量密度)[为可达约最高]*\s*([0-9.]+\s*Wh/kg)", "能量密度"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:探测距离|测距)[为可达约最高]*\s*([0-9.]+\s*m)", "探测距离"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:峰值功率|功率)[可覆盖为达约最高]*\s*([0-9.]+(?:～|-|~)?[0-9.]*\s*kW)", "峰值功率"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:综合续驶里程|纯电续航|续航里程|续航)[超过可达为约]*\s*([0-9.]+\+?\s*km)", "续航里程"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:工况效率|转换效率|峰值效率|效率)[高达大于为约]*\s*([0-9.]+%以上?|[0-9.]+%)", "效率"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:循环寿命)[超过可达为约]*\s*([0-9.]+次)", "循环寿命"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:测距能力)[为可达约最高实现]*\s*([0-9.]+m@[0-9.]+%)", "测距能力"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:角分辨率)[为可达约最高实现]*\s*([0-9.]+°×[0-9.]+°)", "角分辨率"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:充电峰值倍率)[达为约]*\s*([0-9.]+C)", "充电倍率"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:单枪最大支持)[达为约]*\s*([0-9.]+\s*kW)", "单枪最大功率"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:整机噪声小于|噪声小于)\s*([0-9.]+\s*dB)", "噪声"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:支持最高)\s*([0-9.]+V/[0-9.]+A)", "输出规格"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:渗透率)[达超过为约]*\s*([0-9.]+%)", "渗透率"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:减重)[超超过达为约]*\s*([0-9.]+%|1/3)", "减重比例"),
        ]
        for pattern, rel in patterns:
            for m in re.finditer(pattern, sentence):
                subject = normalize_tech_name(m.group(1))
                if is_valid_tech_entity(subject):
                    triples.append({"subject": subject, "relation": rel, "object": m.group(2).replace(" ", ""), "type": "attribute", "source": source_name, "evidence": sentence[:240]})

        relation_patterns = [
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:采用|搭载|使用)([\u4e00-\u9fffA-Za-z0-9\- +.]{2,40})", "采用"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:支持)([\u4e00-\u9fffA-Za-z0-9\- +.]{2,40})", "支持"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:用于|应用于)([\u4e00-\u9fffA-Za-z0-9\- +.]{2,40})", "用于"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:包括|包含)([\u4e00-\u9fffA-Za-z0-9\- +、.]{2,60})", "包含"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:适用于)([\u4e00-\u9fffA-Za-z0-9\- +、.]{2,50})", "适用场景"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:通过了|通过)([\u4e00-\u9fffA-Za-z0-9\- +、.]{2,50})", "通过测试"),
            (r"([\u4e00-\u9fffA-Za-z0-9\- ]{2,30}?)(?:依托)([\u4e00-\u9fffA-Za-z0-9\- +、.]{2,50})", "依托"),
        ]
        for pattern, rel in relation_patterns:
            for m in re.finditer(pattern, sentence):
                subj = normalize_tech_name(m.group(1).strip(" ，,"))
                obj = m.group(2).strip(" ，,")
                if is_valid_tech_entity(subj) and 1 < len(obj) <= 50:
                    triples.append({"subject": subj, "relation": rel, "object": obj, "type": "relation", "source": source_name, "evidence": sentence[:240]})

        for company in TECH_COMPANIES:
            m = re.search(rf"{re.escape(company)}(?:发布|推出|展示|宣布|官宣|升级至|打造了)(?:了|的)?([\u4e00-\u9fffA-Za-z0-9\- +.·“”]+?)(?:，|。|；|采用|搭载|支持|实现|具备|具有|将|于|$)", sentence)
            if m:
                tech = normalize_tech_name(m.group(1).strip("“” "))
                if is_valid_tech_entity(tech) and any(suffix in tech for suffix in ["技术", "系统", "方案", "平台", "产品", "芯片", "电池", "底盘", "雷达", "电机", "电控", "架构", "模型", "座舱"]):
                    triples.append({"subject": tech, "relation": "研发企业", "object": company, "type": "relation", "source": source_name, "evidence": sentence[:240]})
    return triples


def build_technical_entity_graph(triples, source_index=None):
    entities = {}

    def ensure(name, category=None):
        name = normalize_tech_name(name)
        if name not in entities:
            entities[name] = {"type": "entity", "category": category or "Technology", "attributes": {}, "relations": [], "sources": []}
        elif category and entities[name].get("category") in ("Technology", "", None):
            entities[name]["category"] = category
        return entities[name]

    def add_relation(entity, relation, obj, source, evidence):
        values = obj if isinstance(obj, list) else [obj]
        values = [normalize_tech_name(v) for v in values if str(v).strip()]
        if not values:
            return
        for rel in entity["relations"]:
            if rel["relation"] == relation:
                old = rel["object"] if isinstance(rel["object"], list) else [rel["object"]]
                rel["object"] = list(dict.fromkeys(old + values))
                if source and source not in rel.setdefault("sources", []):
                    rel["sources"].append(source)
                if evidence and evidence not in rel.setdefault("evidence", []):
                    rel["evidence"].append(evidence)
                return
        item = {"relation": relation, "object": values[0] if len(values) == 1 else values}
        if source:
            item["sources"] = [source]
        if evidence:
            item["evidence"] = [evidence]
        entity["relations"].append(item)

    seen = set()
    for raw in triples:
        if not all(k in raw for k in ("subject", "relation", "object")):
            continue
        subject = normalize_tech_name(raw["subject"])
        if not is_valid_tech_entity(subject):
            continue
        relation = normalize_tech_relation(raw["relation"])
        obj = raw["object"]
        typ = raw.get("type", "relation")
        source = raw.get("source", "")
        evidence = raw.get("evidence", "")
        key = (subject, relation, json.dumps(obj, ensure_ascii=False, sort_keys=True), typ)
        if key in seen:
            continue
        seen.add(key)

        category = str(obj) if relation == "技术类别" and typ == "attribute" else None
        entity = ensure(subject, category)
        if source and source not in entity["sources"]:
            entity["sources"].append(source)
        if typ == "attribute":
            if relation == "技术类别":
                entity["category"] = str(obj)
            else:
                entity["attributes"][relation] = obj
        else:
            add_relation(entity, relation, obj, source, evidence)
            targets = obj if isinstance(obj, list) else [obj]
            for target in targets:
                target = normalize_tech_name(target)
                if target and len(target) <= 40 and target != subject:
                    ensure(target)

    total_attr = sum(len(e["attributes"]) for e in entities.values())
    total_rel = sum(len(e["relations"]) for e in entities.values())
    return {
        "meta": {
            "source": "technical_seed_records + rule_ie",
            "source_index": str(source_index) if source_index else "",
            "entities": len(entities),
            "total_attributes": total_attr,
            "total_relations": total_rel,
            "focus": "automotive technology graph",
        },
        "entities": dict(sorted(entities.items(), key=lambda kv: kv[0])),
    }


def run_technical_mode(args):
    input_files = [Path(p).resolve() for p in (args.input or [])]
    for path in input_files:
        if not path.exists():
            raise FileNotFoundError(path)

    triples = list(TECH_SEED_TRIPLES)
    for path in input_files:
        print(f"[technical] reading {path}")
        text = read_technical_text(path)
        triples.extend(extract_technical_triples_by_rules(text, path.name))

    output_path = Path(args.output).resolve() if args.output else TECH_DEFAULT_OUTPUT
    output_path.parent.mkdir(parents=True, exist_ok=True)
    graph = build_technical_entity_graph(triples, args.source_index or TECH_DEFAULT_SOURCE_INDEX)
    output_path.write_text(json.dumps(graph, ensure_ascii=False, indent=2), encoding="utf-8")

    print(f"wrote: {output_path}")
    print(f"entities: {graph['meta']['entities']}")
    print(f"attributes: {graph['meta']['total_attributes']}")
    print(f"relations: {graph['meta']['total_relations']}")


def parse_cli_args():
    parser = argparse.ArgumentParser(description="Extract offline technology triples for the automotive QA graph.")
    parser.add_argument("--mode", choices=["technical"], default="technical", help="technical builds the automotive technology graph without calling LLMs.")
    parser.add_argument("--input", nargs="*", default=[], help="Input files for technical mode.")
    parser.add_argument("--output", default=None, help="Output JSON path.")
    parser.add_argument("--source-index", default=None, help="knowledge_sources.json path for technical mode.")
    return parser.parse_args()


def main():
    args = parse_cli_args()
    run_technical_mode(args)


if __name__ == "__main__":
    main()
